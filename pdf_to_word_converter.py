#!/usr/bin/env python3
# -*- coding: utf-8 -*-


import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import os
import tempfile


def pdf_to_word_pymupdf(pdf_path, output_path="output.docx", zoom=2.5):
    
    if not os.path.exists(pdf_path):
        print(f"❌ خطأ: الملف '{pdf_path}' غير موجود")
        return

    try:
        doc_pdf = fitz.open(pdf_path)
    except Exception as e:
        print(f"❌ خطأ في فتح ملف PDF: {e}")
        return

    total_pages = doc_pdf.page_count
    if total_pages <= 1:
        print("❌ الملف يحتوي على صفحة واحدة فقط...")
        doc_pdf.close()
        return

    print(f"� عدد الصفحات في PDF: {total_pages}")
    print("� سيتم استثناء الصفحة الأولى ومعالجة الباقي...")

    temp_dir = tempfile.mkdtemp(prefix="pdf_images_")
    image_paths = []

    for page_index in range(1, total_pages):
        try:
            page = doc_pdf.load_page(page_index)
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)

            img_path = os.path.join(temp_dir, f"page_{page_index + 1}.png")
            pix.save(img_path)
            image_paths.append((img_path, page_index + 1))
            print(f"✅ تم إنشاء صورة للصفحة {page_index + 1}")
        except Exception as e:
            print(f"⚠️ خطأ في تحويل الصفحة {page_index + 1}: {e}")

    doc_pdf.close()

    if not image_paths:
        print("❌ لم يتم إنشاء أي صور")
        return

    print("� جاري إنشاء ملف Word...")

    doc = Document()
    section = doc.sections[0]

    # صفحة أفقية بدون هوامش كبيرة
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.1)  # هامش سفلي صغير جداً

    # صورتان في كل صفحة
    for i in range(0, len(image_paths), 2):
        group = image_paths[i:i + 2]

        table = doc.add_table(rows=1, cols=len(group))
        table.autofit = False
        table.allow_autofit = False

        # عرض مناسب لكل عمود
        col_width = Inches(5.0)
        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for j, (img_path, page_num) in enumerate(group):
            cell = table.rows[0].cells[j]
            cell.text = ""
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

            try:
                run = paragraph.add_run()
                # صور مصغرة لإظهار كل المحتوى
                run.add_picture(img_path, height=Inches(5.2))
            except Exception as e:
                print(f"⚠️ خطأ في إضافة صورة الصفحة {page_num}: {e}")
                continue

            # رقم الصفحة مع مسافة قليلة جداً
            page_para = cell.add_paragraph()
            page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            page_para.paragraph_format.space_before = Pt(1)
            page_para.paragraph_format.space_after = Pt(0)
            
            page_run = page_para.add_run(f"ص. {page_num}")
            page_run.font.size = Pt(9)
            page_run.font.color.rgb = RGBColor(120, 120, 120)

        # فاصل صفحة
        if i + 2 < len(image_paths):
            doc.add_page_break()

    try:
        doc.save(output_path)
        print(f"✅ تم إنشاء ملف Word بنجاح: {output_path}")
        print(f"� عدد الصفحات المعالجة: {len(image_paths)}")
    except Exception as e:
        print(f"❌ خطأ في حفظ ملف Word: {e}")

    for img_path, _ in image_paths:
        try:
            os.remove(img_path)
        except:
            pass
    try:
        os.rmdir(temp_dir)
    except:
        pass


def main():
    print("=" * 50)
    print("� محول PDF إلى Word - نسخة محسّنة نهائية")
    print("=" * 50)

    pdf_file = input("� أدخل مسار ملف PDF: ").strip()
    output_file = input("� اسم ملف Word (افتراضي 'output.docx'): ").strip()
    if not output_file:
        output_file = "output.docx"

    zoom_inp = input("� عامل التكبير للوضوح (افتراضي 2.5): ").strip()
    try:
        zoom = float(zoom_inp) if zoom_inp else 2.5
    except ValueError:
        zoom = 2.5

    pdf_to_word_pymupdf(pdf_file, output_file, zoom)


if __name__ == "__main__":
    main()
